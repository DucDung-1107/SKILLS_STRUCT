#!/usr/bin/env python3
"""
📁 File Processing Utilities
Các hàm tiện ích cho xử lý file và I/O operations
"""

import os
import json
import csv
import pandas as pd
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import logging
import mimetypes
import zipfile
import tempfile
from datetime import datetime

# PDF and document processing
try:
    import PyMuPDF as fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyMuPDF not available. PDF processing will be limited.")

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing will be limited.")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

logger = logging.getLogger(__name__)

# Supported file types
SUPPORTED_EXTENSIONS = {
    'documents': ['.pdf', '.doc', '.docx', '.txt', '.rtf'],
    'data': ['.json', '.csv', '.xlsx', '.xls'],
    'images': ['.jpg', '.jpeg', '.png', '.gif', '.bmp'],
    'archives': ['.zip', '.rar', '.7z']
}

def read_json_file(file_path: str) -> Dict[str, Any]:
    """
    Đọc file JSON
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in {file_path}: {e}")
        return {}
    except Exception as e:
        logger.error(f"Error reading JSON file {file_path}: {e}")
        return {}

def write_json_file(data: Dict[str, Any], file_path: str, indent: int = 2) -> bool:
    """
    Ghi file JSON
    """
    try:
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=indent, ensure_ascii=False)
        return True
    except Exception as e:
        logger.error(f"Error writing JSON file {file_path}: {e}")
        return False

def read_csv_file(file_path: str, **kwargs) -> pd.DataFrame:
    """
    Đọc file CSV
    """
    try:
        return pd.read_csv(file_path, **kwargs)
    except FileNotFoundError:
        logger.error(f"CSV file not found: {file_path}")
        return pd.DataFrame()
    except Exception as e:
        logger.error(f"Error reading CSV file {file_path}: {e}")
        return pd.DataFrame()

def write_csv_file(data: Union[pd.DataFrame, List[Dict]], file_path: str, **kwargs) -> bool:
    """
    Ghi file CSV
    """
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        if isinstance(data, pd.DataFrame):
            data.to_csv(file_path, index=False, **kwargs)
        else:
            # Convert list of dicts to DataFrame
            df = pd.DataFrame(data)
            df.to_csv(file_path, index=False, **kwargs)
        return True
    except Exception as e:
        logger.error(f"Error writing CSV file {file_path}: {e}")
        return False

def read_excel_file(file_path: str, sheet_name: str = None) -> pd.DataFrame:
    """
    Đọc file Excel
    """
    try:
        return pd.read_excel(file_path, sheet_name=sheet_name)
    except FileNotFoundError:
        logger.error(f"Excel file not found: {file_path}")
        return pd.DataFrame()
    except Exception as e:
        logger.error(f"Error reading Excel file {file_path}: {e}")
        return pd.DataFrame()

def write_excel_file(data: Union[pd.DataFrame, Dict[str, pd.DataFrame]], 
                    file_path: str) -> bool:
    """
    Ghi file Excel
    """
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        if isinstance(data, pd.DataFrame):
            data.to_excel(file_path, index=False)
        elif isinstance(data, dict):
            # Multiple sheets
            with pd.ExcelWriter(file_path) as writer:
                for sheet_name, df in data.items():
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
        return True
    except Exception as e:
        logger.error(f"Error writing Excel file {file_path}: {e}")
        return False

def extract_pdf_text(file_path: str) -> str:
    """
    Trích xuất text từ PDF
    """
    if not PDF_AVAILABLE:
        logger.warning("PyMuPDF not available for PDF processing")
        return ""
    
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        doc.close()
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text from {file_path}: {e}")
        return ""

def extract_doc_text(file_path: str) -> str:
    """
    Trích xuất text từ DOC/DOCX
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.docx':
        return extract_docx_text(file_path)
    elif file_ext == '.doc':
        # For .doc files, we need python-docx or docx2txt
        if DOCX2TXT_AVAILABLE:
            try:
                return docx2txt.process(file_path)
            except Exception as e:
                logger.error(f"Error extracting DOC text from {file_path}: {e}")
                return ""
        else:
            logger.warning("docx2txt not available for DOC processing")
            return ""
    else:
        logger.warning(f"Unsupported document format: {file_ext}")
        return ""

def extract_docx_text(file_path: str) -> str:
    """
    Trích xuất text từ DOCX
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available for DOCX processing")
        return ""
    
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX text from {file_path}: {e}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """
    Trích xuất text từ file dựa trên extension
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return ""
    
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return extract_pdf_text(file_path)
    elif file_ext in ['.doc', '.docx']:
        return extract_doc_text(file_path)
    elif file_ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    else:
        logger.warning(f"Unsupported file type for text extraction: {file_ext}")
        return ""

def validate_file_type(file_path: str, allowed_extensions: List[str] = None) -> bool:
    """
    Validate file type
    """
    if not os.path.exists(file_path):
        return False
    
    if allowed_extensions is None:
        # Default allowed extensions for resume files
        allowed_extensions = ['.pdf', '.doc', '.docx', '.txt']
    
    file_ext = Path(file_path).suffix.lower()
    return file_ext in allowed_extensions

def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Lấy thông tin file
    """
    if not os.path.exists(file_path):
        return {}
    
    stat = os.stat(file_path)
    file_path_obj = Path(file_path)
    
    # Get MIME type
    mime_type, _ = mimetypes.guess_type(file_path)
    
    return {
        'name': file_path_obj.name,
        'extension': file_path_obj.suffix.lower(),
        'size': stat.st_size,
        'size_mb': round(stat.st_size / (1024 * 1024), 2),
        'created': datetime.fromtimestamp(stat.st_ctime),
        'modified': datetime.fromtimestamp(stat.st_mtime),
        'mime_type': mime_type,
        'is_file': os.path.isfile(file_path),
        'is_dir': os.path.isdir(file_path)
    }

def create_backup(file_path: str, backup_dir: str = None) -> str:
    """
    Tạo backup file
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if backup_dir is None:
        backup_dir = os.path.dirname(file_path)
    
    os.makedirs(backup_dir, exist_ok=True)
    
    file_path_obj = Path(file_path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"{file_path_obj.stem}_backup_{timestamp}{file_path_obj.suffix}"
    backup_path = os.path.join(backup_dir, backup_name)
    
    import shutil
    shutil.copy2(file_path, backup_path)
    
    return backup_path

def extract_zip_file(zip_path: str, extract_to: str = None) -> List[str]:
    """
    Giải nén file ZIP và trả về list các file đã extract
    """
    if extract_to is None:
        extract_to = os.path.dirname(zip_path)
    
    extracted_files = []
    
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            # Get list of files in zip
            file_list = zip_ref.namelist()
            
            for file_name in file_list:
                # Skip directories
                if not file_name.endswith('/'):
                    # Extract file
                    zip_ref.extract(file_name, extract_to)
                    extracted_file_path = os.path.join(extract_to, file_name)
                    extracted_files.append(extracted_file_path)
    
    except Exception as e:
        logger.error(f"Error extracting ZIP file {zip_path}: {e}")
    
    return extracted_files

def create_temp_file(content: str = None, suffix: str = '.txt') -> str:
    """
    Tạo temporary file
    """
    temp_file = tempfile.NamedTemporaryFile(mode='w', suffix=suffix, delete=False)
    
    if content:
        temp_file.write(content)
    
    temp_file.close()
    return temp_file.name

def clean_temp_files(temp_dir: str = None, older_than_hours: int = 24):
    """
    Xóa temporary files cũ
    """
    if temp_dir is None:
        temp_dir = tempfile.gettempdir()
    
    current_time = datetime.now()
    cutoff_time = current_time.timestamp() - (older_than_hours * 3600)
    
    deleted_count = 0
    
    try:
        for file_path in Path(temp_dir).glob('*'):
            if file_path.is_file():
                if file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        deleted_count += 1
                    except Exception as e:
                        logger.warning(f"Could not delete temp file {file_path}: {e}")
    except Exception as e:
        logger.error(f"Error cleaning temp files: {e}")
    
    return deleted_count

def safe_filename(filename: str, max_length: int = 255) -> str:
    """
    Tạo safe filename (remove invalid characters)
    """
    # Remove/replace invalid characters
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    
    # Remove leading/trailing spaces and dots
    filename = filename.strip(' .')
    
    # Truncate if too long
    if len(filename) > max_length:
        name_part, ext_part = os.path.splitext(filename)
        max_name_len = max_length - len(ext_part)
        filename = name_part[:max_name_len] + ext_part
    
    return filename

def get_unique_filename(file_path: str) -> str:
    """
    Tạo unique filename nếu file đã tồn tại
    """
    if not os.path.exists(file_path):
        return file_path
    
    file_path_obj = Path(file_path)
    directory = file_path_obj.parent
    name = file_path_obj.stem
    extension = file_path_obj.suffix
    
    counter = 1
    while True:
        new_name = f"{name}_{counter}{extension}"
        new_path = directory / new_name
        if not new_path.exists():
            return str(new_path)
        counter += 1

def batch_process_files(file_paths: List[str], process_func, 
                       batch_size: int = 10, **kwargs) -> List[Any]:
    """
    Xử lý batch files
    """
    results = []
    
    for i in range(0, len(file_paths), batch_size):
        batch = file_paths[i:i + batch_size]
        
        for file_path in batch:
            try:
                result = process_func(file_path, **kwargs)
                results.append({
                    'file_path': file_path,
                    'success': True,
                    'result': result
                })
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {e}")
                results.append({
                    'file_path': file_path,
                    'success': False,
                    'error': str(e)
                })
    
    return results

def monitor_file_changes(file_path: str, callback_func, interval: int = 1):
    """
    Monitor file changes and call callback when changed
    """
    import time
    
    if not os.path.exists(file_path):
        logger.error(f"File not found for monitoring: {file_path}")
        return
    
    last_modified = os.path.getmtime(file_path)
    
    while True:
        try:
            current_modified = os.path.getmtime(file_path)
            if current_modified != last_modified:
                callback_func(file_path)
                last_modified = current_modified
            
            time.sleep(interval)
        except KeyboardInterrupt:
            break
        except Exception as e:
            logger.error(f"Error monitoring file {file_path}: {e}")
            break

def calculate_file_hash(file_path: str, algorithm: str = 'md5') -> str:
    """
    Tính hash của file
    """
    import hashlib
    
    hash_func = getattr(hashlib, algorithm.lower(), None)
    if not hash_func:
        raise ValueError(f"Unsupported hash algorithm: {algorithm}")
    
    hasher = hash_func()
    
    try:
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception as e:
        logger.error(f"Error calculating hash for {file_path}: {e}")
        return ""
